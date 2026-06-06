"""
# reporter.py
# ─────────────────────────────────────────────────────────────
# Generates the Word (.docx) assessment report.
# Takes the incident details collected by main.py and the
# Assessment result from assessor.py, then builds a
# professionally formatted document suitable for attaching
# to an organisation's incident register.
#
# Uses python-docx for all Word document construction.
# ─────────────────────────────────────────────────────────────
"""


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from assessor import Assessment


# Maps harm_score (0–6) to a risk label and display colour.
# Used to render the severity rating section of the report.
# Colours follow a traffic-light convention: green → amber → red.

RISK_LABELS = {
    0: ("Low", RGBColor(0x0F, 0x6E, 0x56)),
    1: ("Low", RGBColor(0x0F, 0x6E, 0x56)),
    2: ("Medium", RGBColor(0xBA, 0x75, 0x17)),
    3: ("Medium", RGBColor(0xBA, 0x75, 0x17)),
    4: ("High", RGBColor(0xD8, 0x5A, 0x30)),
    5: ("Critical", RGBColor(0xA3, 0x2D, 0x2D)),
    6: ("Critical", RGBColor(0xA3, 0x2D, 0x2D)),
}


def _heading(doc, text, level=1):

    """
    Add a styled heading to the document.
    Applies the brand blue colour to all headings for visual consistency.
    python-docx requires accessing run[0] after add_heading to style the text.
    """
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)


def _gate_row(table, gate_num, label, result):

    """
    Append one row to the gate analysis table.
    Each row covers one gate: its number, the test name,
    pass/fail status (colour-coded), and the reasoning string
    from the GateResult object.
    """
    row = table.add_row()
    row.cells[0].text = f"Gate {gate_num}"
    row.cells[1].text = label
    status = "PASSED" if result.passed else "NOT MET"
    row.cells[2].text = status
    row.cells[2].paragraphs[0].runs[0].font.color.rgb = (
        RGBColor(0x0F, 0x6E, 0x56) if result.passed else RGBColor(0xA3, 0x2D, 0x2D)
    )
    row.cells[3].text = result.reason


def generate_report(details: dict, assessment: Assessment, output_path: str):
    
    """
    Build and save the full assessment report as a .docx file.

    Parameters:
        details     — dict of free-text incident fields from main.py
        assessment  — Assessment dataclass from assessor.assess()
        output_path — file path string where the .docx will be saved
    """
    doc = Document()

    # Title
    title = doc.add_heading("NDB Scheme — Data Breach Assessment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Legal authority citation - establishes the document's basis
    doc.add_paragraph(
        "Prepared under Part IIIC of the Privacy Act 1988 (Cth) "
        "and the OAIC Notifiable Data Breaches scheme guidelines."
    ).italic = True

    # Incident details
    _heading(doc, "1. Incident details")

    #Two-column table: field name (bold) | value
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    fields = [
        ("Organisation", details["organisation"]),
        ("Date discovered", details["incident_date"]),
        ("Assessor", details["assessor_name"]),
        ("Report date", date.today().strftime("%d/%m/%Y")),
        ("Individuals affected (est.)", details["individuals"]),
        ("Data types involved", details["data_types"]),
    ]
    for label, value in fields:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = value

    doc.add_paragraph()

    # Section 2: Incident description
    _heading(doc, "2. Incident description")
    doc.add_paragraph(details["incident_desc"])

    # Section 3: Gate analysis
    _heading(doc, "3. NDB eligibility assessment")
    doc.add_paragraph(
        "The following gates were applied in accordance with s26WE of the Privacy Act 1988 (Cth)."
    )

    # Four-column table: Gate | Test | Result | Reasoning
    gate_table = doc.add_table(rows=1, cols=4)
    gate_table.style = "Table Grid"

    #Header row - bold labels
    hdr = gate_table.rows[0].cells
    for i, h in enumerate(["Gate", "Test", "Result", "Reasoning"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    # Add a row for each gate that was evaluated.
    # gate2 and gate3 may be None if the assessment exited early.
    _gate_row(gate_table, 1, "Eligible data breach", assessment.gate1)
    if assessment.gate2:
        _gate_row(gate_table, 2, "Likely serious harm", assessment.gate2)
    if assessment.gate3:
        _gate_row(gate_table, 3, "No exemption applies", assessment.gate3)

    doc.add_paragraph()

    # Harm score
    # ── Section 4: Harm severity rating ───────────────────────
    # Look up the label and colour for the harm score (0–6).

    risk_label, risk_color = RISK_LABELS[assessment.harm_score]
    _heading(doc, "4. Harm severity rating")
    p = doc.add_paragraph()
    run = p.add_run(f"{risk_label}  ({assessment.harm_score}/6 harm factors present)")
    run.font.bold = True
    run.font.color.rgb = risk_color
    run.font.size = Pt(13)

    # Section 5: Outcome
    _heading(doc, "5. Outcome and required action")
    outcome_para = doc.add_paragraph()
    outcome_run = outcome_para.add_run(assessment.outcome_summary)
    outcome_run.font.bold = True
    if assessment.notifiable:
        #Red verdict - notification is legally required.
        outcome_run.font.color.rgb = RGBColor(0xA3, 0x2D, 0x2D)

        # Numbered list of mandatory actions under the NDB scheme
        doc.add_paragraph(
            "Required actions:\n"
            "1. Notify the Office of the Australian Information Commissioner (OAIC) "
            "via the online NDB notification form at oaic.gov.au.\n"
            "2. Notify all affected individuals as soon as practicable.\n"
            "3. Notification must be completed within 30 days of becoming aware "
            "of the eligible data breach (s26WK).\n"
            "4. Retain this assessment document as part of your incident record."
        )
    else:
        outcome_run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
        doc.add_paragraph(
            "No notification to the OAIC or individuals is required at this time. "
            "This assessment document should be retained on file in accordance with "
            "your organisation's records management obligations."
        )

    # Section 6: Draft notification (only if notifiable)
    # Only generated when all three gates pass.
    # Provides a pre-filled template for the assessor to complete
    # and submit to the OAIC — bracketed fields need manual completion.

    if assessment.notifiable:
        _heading(doc, "6. Draft OAIC notification statement")
        doc.add_paragraph(
            "The following is a draft notification for review by your legal or privacy team "
            "before submission to the OAIC. This is not legal advice."
        ).italic = True
        doc.add_paragraph(
            f"{details['organisation']} hereby notifies the Office of the Australian "
            f"Information Commissioner of an eligible data breach discovered on "
            f"{details['incident_date']}.\n\n"
            f"Nature of the breach: {details['incident_desc']}\n\n"
            f"Types of personal information involved: {details['data_types']}\n\n"
            f"Estimated number of individuals affected: {details['individuals']}\n\n"
            f"Steps taken in response: [TO BE COMPLETED — describe containment, "
            f"investigation, and remediation actions taken.]\n\n"
            f"Recommended steps for affected individuals: [TO BE COMPLETED — e.g. "
            f"change passwords, monitor financial accounts, place credit alerts.]"
        )

    # Footer disclaimer
    # Clearly marks the document as a tool output, not legal advice.
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "DISCLAIMER: This report is generated by an automated assessment tool for "
        "educational and portfolio purposes. It does not constitute legal advice. "
        "Consult a qualified privacy lawyer or the OAIC guidelines before making "
        "any notification decisions."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x87, 0x80)

    # Save the completed document to the specified path.
    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")